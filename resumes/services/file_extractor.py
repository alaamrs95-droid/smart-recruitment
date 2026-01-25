import os
from pathlib import Path
import pdfplumber # pip install pdfplumber
from docx import Document  # pip install python-docx
from pypdf import PdfReader

def extract_text(file_path: str) -> str:
    """استخراج النص من الملفات بدقة عالية"""
    suffix = Path(file_path).suffix.lower()
    
    try:
        if suffix == '.pdf':
            return extract_text_from_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        elif suffix == '.txt':
            return extract_text_from_txt(file_path)
        else:
            return "Text extraction not implemented for this file type"
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def extract_text_from_pdf(file_path: str) -> str:
    """استخراج نص من PDF باستخدام pypdf"""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception:
        text = ""  # fallback فارغ
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """استخراج نص من DOCX"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        # استخراج من tables كمان
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
    except:
        text = ""
    return text.strip()

def extract_text_from_txt(file_path: str) -> str:
    """استخراج النص من ملف نصي (كما هو مع تحسين)"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except:
        try:
            with open(file_path, 'r', encoding='cp1256', errors='ignore') as file:
                return file.read()
        except:
            return ""